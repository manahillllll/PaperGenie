import os
os.environ['TF_ENABLE_ONEDNN_OPTS'] = '0'
import requests
import xml.etree.ElementTree as ET
import fitz  # PyMuPDF
from docx import Document
import re
from transformers import pipeline

# ========== Load Model ==========

# Load a pre-trained BART or T5 summarizer from Hugging Face
summarizer = pipeline("summarization", model="facebook/bart-large-cnn")

# ========== Core Functions ==========

def search_arxiv(query="AI medical diagnosis", max_results=3):
    base_url = "http://export.arxiv.org/api/query?"
    search_query = f"search_query=all:{query.replace(' ', '+')}&start=0&max_results={max_results}"
    url = base_url + search_query
    response = requests.get(url)

    root = ET.fromstring(response.content)
    ns = {'atom': 'http://www.w3.org/2005/Atom'}

    papers = []
    for entry in root.findall('atom:entry', ns):
        title = entry.find('atom:title', ns).text.strip()
        summary = entry.find('atom:summary', ns).text.strip()
        link = entry.find('atom:id', ns).text.strip()
        authors = [a.find('atom:name', ns).text for a in entry.findall('atom:author', ns)]
        published = entry.find('atom:published', ns).text.split("T")[0]
        arxiv_id = link.split('/')[-1]

        citation = f"@article{{{authors[0].split()[-1].lower()}{published[:4]}," \
                   f"\n  title={{\"{title}\"}}," \
                   f"\n  author={{\"{' and '.join(authors)}\"}}," \
                   f"\n  journal={{arXiv preprint arXiv:{arxiv_id}}}," \
                   f"\n  year={{\"{published[:4]}\"}}" \
                   f"\n}}"

        papers.append({
            'title': title,
            'summary': summary,
            'link': link,
            'citation': citation,
            'id': arxiv_id
        })
    return papers

def download_pdf(arxiv_id, save_dir="pdfs"):
    os.makedirs(save_dir, exist_ok=True)
    pdf_url = f"https://arxiv.org/pdf/{arxiv_id}.pdf"
    path = os.path.join(save_dir, f"{arxiv_id}.pdf")
    r = requests.get(pdf_url)
    if r.status_code == 200:
        with open(path, "wb") as f:
            f.write(r.content)
        return path
    return None

def extract_text_from_pdf(pdf_path):
    doc = fitz.open(pdf_path)
    text = "\n".join([page.get_text() for page in doc])
    return text

# Updated extractive summary using HuggingFace's BART model
def extractive_summary(text, max_length=200):
    # Truncate the text to ~4000 characters (approx. 1024 tokens)
    truncated_text = text[:4000]
    if len(truncated_text.strip()) < 100:
        return "❗ Not enough content to summarize."

    # Run the summarizer on the truncated text
    summary = summarizer(truncated_text, max_length=max_length, min_length=50, do_sample=False)
    return summary[0]['summary_text']
# ========== DOCX Writer ==========

def save_doc(papers, filename="Summarized_Papers.docx"):
    doc = Document()
    doc.add_heading("AI Research Summaries", 0)

    for paper in papers:
        doc.add_heading(paper['title'], level=1)
        doc.add_paragraph(f" Link: {paper['link']}", style='Intense Quote')
        doc.add_paragraph(" Summary:", style='Heading 2')
        
        # Clean the summary text before adding it
        def clean_text(text):
            return re.sub(r'[\x00-\x1F\x7F]', '', text)  # Remove control characters and delete characters

        doc.add_paragraph(clean_text(paper['summary']))
        doc.add_paragraph(" Citation (BibTeX):", style='Heading 2')
        doc.add_paragraph(paper['citation'])
        doc.add_paragraph("\n")

    doc.save(filename)
    return filename

# ========== Web UI Logic ==========

def run_pipeline(query, max_results):
    papers = search_arxiv(query, max_results=max_results)
    final_results = []

    for paper in papers:
        pdf_path = download_pdf(paper['id'])
        if not pdf_path:
            paper['summary'] = "❌ PDF not found"
        else:
            try:
                text = extract_text_from_pdf(pdf_path)
                # Summarize the extracted text
                paper['summary'] = extractive_summary(text)
            except Exception as e:
                paper['summary'] = f"⚠️ Error: {str(e)}"
        final_results.append(paper)

    doc_file = save_doc(final_results)
    return "\n\n".join([f"📌 {p['title']}\n🧠 Summary: {p['summary']}\n🔗 {p['link']}" for p in final_results]), doc_file

# ========== Gradio App ==========

import gradio as gr

with gr.Blocks(title="PaperGenie 📚") as demo:
    gr.Markdown("# 🔍 PaperGenie: Your AI Research Assistant")
    gr.Markdown("Enter a research topic below. PaperGenie will fetch, summarize, and generate citations for top arXiv papers. ✨")

    with gr.Row():
        query = gr.Textbox(label="Enter Research Query", placeholder="e.g. AI in medical diagnosis")
        num_results = gr.Slider(minimum=1, maximum=10, step=1, value=3, label="Number of Papers")

    btn = gr.Button("Generate Summaries")
    output_text = gr.Textbox(label="Summaries", lines=20)
    output_docx = gr.File(label="Download .docx Report")

    btn.click(fn=run_pipeline, inputs=[query, num_results], outputs=[output_text, output_docx])

demo.launch()
